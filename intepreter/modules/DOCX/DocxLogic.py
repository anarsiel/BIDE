import os
from io import BytesIO
from urllib.request import urlopen

import PIL.Image
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

# from modules._interfaces.CommonLogic import CommonLogic
from intepreter.modules._interfaces.CommonLogic import CommonLogic


class DocxLogic:
    __document = None
    __paragraph = None

    @staticmethod
    def open_new_document():
        if DocxLogic.__document:
            raise CommonLogic.RunTimeError(
                f'Before opening new Document, close the other.'
            )

        DocxLogic.__document = Document()
        DocxLogic.__paragraph = DocxLogic.__document.add_paragraph()

    @staticmethod
    def save_document_as(document_name, hard=False):
        if not DocxLogic.__document:
            raise CommonLogic.RunTimeError(
                f'Before saving Document, open some.'
            )

        if not hard and os.path.exists(document_name):
            raise CommonLogic.RunTimeError(
                f'Document name `{document_name}` is already taken.'
            )

        DocxLogic.__document.save(document_name)

    @staticmethod
    def save_hard_document_as(document_name):
       DocxLogic.save_document_as(document_name, hard=True)

    @staticmethod
    def new_page():
        r = DocxLogic.__paragraph.add_run()
        r.add_break(WD_BREAK.PAGE)

    @staticmethod
    def new_line():
        r = DocxLogic.__paragraph.add_run()
        r.add_break(WD_BREAK.LINE)

    @staticmethod
    def new_paragraph():
        DocxLogic.__paragraph = DocxLogic.__document.add_paragraph()

    @staticmethod
    def add_text(text):
        r = DocxLogic.__paragraph.add_run()
        r.add_text(text)

    @staticmethod
    def add_title(text):
        r = DocxLogic.__paragraph.add_run()
        r.font.size = Pt(72)
        r.font.bold = True
        r.add_text(text)

    @staticmethod
    def add_image(filename):
        r = DocxLogic.__paragraph.add_run()
        r.add_picture(filename)

    @staticmethod
    def add_image_by_url(url):
        try:
            response = urlopen(url)
            img = PIL.Image.open(BytesIO(response.read()))

            img_bytes = BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)

            DocxLogic.__document.add_picture(img_bytes)
        except:
            CommonLogic.RunTimeError(f"Wrong url: {url}")

    @staticmethod
    def reboot():
        DocxLogic.__document = None
        DocxLogic.__paragraph = None


class Docx:
    __info = [
        ['open_new_document', DocxLogic.open_new_document, [], None, None],
        ['save_as', DocxLogic.save_document_as, [str], None, None],
        ['save_hard_as', DocxLogic.save_hard_document_as, [str], None, None],
        ['new_page', DocxLogic.new_page, [], None, None],
        ['new_line', DocxLogic.new_line, [], None, None],
        ['new_paragraph', DocxLogic.new_paragraph, [], None, None],
        ['add_text', DocxLogic.add_text, [str], None, None],
        ['add_title', DocxLogic.add_title, [str], None, None],
        ['add_image', DocxLogic.add_image, [str], None, None],
        ['add_image_by_url', DocxLogic.add_image_by_url, [str], None, None],
    ]

    @staticmethod
    def get_info():
        return Docx.__info

    @staticmethod
    def reboot():
        DocxLogic.reboot()